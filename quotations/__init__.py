from docx import Document
from typing import Dict

from jinja2 import Template


def template_render(template, replacements):
    return Template(template).render(replacements)


def generate_quote_from_template(template: Document, replacements: Dict):
    doc = template

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = template_render(run.text, replacements)

        paragraph.text = template_render(paragraph.text, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = template_render(paragraph.text, replacements)

    return doc
