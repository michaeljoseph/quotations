"""
TODO:
- [x] cli: quotes.py <path-to-template> <path-to-data>
- [] docx: header / footer content
"""
import csv
import sys
import uuid

from docx import Document
from jinja2 import Template


def template_render(template, replacements):
    return Template(template).render(replacements)


def main(template_document, replacements):
    doc = Document(template_document)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = template_render(run.text, replacements)

        paragraph.text = template_render(paragraph.text, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = template_render(paragraph.text, replacements)

    output_filename = 'generated-proposal-{}.docx'.format(uuid.uuid4().hex)
    print(output_filename)
    doc.save(output_filename)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('usage: quotes.py <path-to-template> <path-to-data>')
        sys.exit(1)

    _, template_path, replacements_path = sys.argv

    replacements = {
        row['variable']: row['value'] for row in csv.DictReader(open(replacements_path))
    }
    main(template_path, replacements)
