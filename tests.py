import pytest
from docx import Document

from quotations import generate_quote_from_template


@pytest.fixture
def template_doc():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        '{{this_is}} the first item in an unordered list', style='ListBullet'
    )
    document.add_paragraph('first item in ordered list', style='ListNumber')
    document.add_paragraph('these are just some {{data}}')
    return document


@pytest.fixture
def template_values():
    return {'this_is': 'just my test', 'data': 'of simple things'}


def doc_to_text(doc: Document):
    text = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text.append(run.text)

        text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text.append(paragraph.text)

    return text


def test_quote_replacement(template_doc, template_values):
    quote_text = doc_to_text(
        generate_quote_from_template(template_doc, template_values)
    )

    assert any(['first item in ordered' in text for text in quote_text])

    for variable, value in template_values.items():
        assert variable not in quote_text
        print([value in text for text in quote_text])
        assert any([value in text for text in quote_text])
